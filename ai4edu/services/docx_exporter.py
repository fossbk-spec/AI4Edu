import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from ai4edu.models.lesson_plan_2345 import LessonPlan2345

def set_cell_background(cell, fill_hex):
    """Đặt màu nền cho cell trong Word table"""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def create_lesson_plan_docx(plan: LessonPlan2345) -> io.BytesIO:
    """
    Xuất Kế hoạch bài dạy chuẩn 5 cột (Công văn 2345) ra file Microsoft Word (.docx) đẹp mắt.
    """
    doc = Document()
    
    # Thiết lập lề trang
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Tiêu đề trường & tên kế hoạch bài dạy
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_school = header_p.add_run(f"{plan.school_name.upper()}\n")
    run_school.font.size = Pt(12)
    run_school.font.bold = True
    run_school.font.color.rgb = RGBColor(11, 87, 208)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("KẾ HOẠCH BÀI DẠY (CÔNG VĂN 2345/BGDĐT-GDTH)\n")
    run_title.font.size = Pt(14)
    run_title.font.bold = True
    
    run_lesson = title_p.add_run(f"BÀI: {plan.lesson_title.upper()}\n")
    run_lesson.font.size = Pt(13)
    run_lesson.font.bold = True
    run_lesson.font.color.rgb = RGBColor(194, 65, 12)

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run(f"Môn: {plan.subject} | Khối lớp: {plan.grade} | Thời lượng: {plan.duration_periods} tiết")
    meta_run.font.italic = True
    meta_run.font.size = Pt(11)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # I. YÊU CẦU CẦN ĐẠT
    h1 = doc.add_heading("I. YÊU CẦU CẦN ĐẠT (MỤC TIÊU BÀI HỌC)", level=2)
    for req in plan.required_competencies:
        p = doc.add_paragraph(req, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    # II. ĐỒ DÙNG DẠY HỌC
    doc.add_heading("II. ĐỒ DÙNG DẠY HỌC & HỌC LIỆU SỐ", level=2)
    for eq in plan.teaching_equipment:
        p = doc.add_paragraph(eq, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    # III. CÁC HOẠT ĐỘNG DẠY HỌC CHỦ YẾU
    doc.add_heading("III. CÁC HOẠT ĐỘNG DẠY HỌC CHỦ YẾU (BẢNG 5 CỘT CV 2345)", level=2)

    for i, act in enumerate(plan.activities, 1):
        doc.add_heading(f"Hoạt động {i}: {act.activity_name}", level=3)
        
        # Bảng chi tiết 5 cột
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # Header của bảng
        headers = ["Hoạt động học", "Mục tiêu", "Nội dung", "Sản phẩm", "Tổ chức thực hiện (4 bước)"]
        hdr_cells = table.rows[0].cells
        for idx, text in enumerate(headers):
            hdr_cells[idx].text = text
            set_cell_background(hdr_cells[idx], "0B57D0")
            for paragraph in hdr_cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(10)

        # Hàng dữ liệu
        row_cells = table.add_row().cells
        row_cells[0].text = act.activity_name
        row_cells[1].text = act.objective
        row_cells[2].text = act.content
        row_cells[3].text = act.product

        # Cột 5: 4 bước tổ chức
        steps_text = []
        for step in act.implementation_steps:
            steps_text.append(f"Bước {step.step_number} ({step.step_name}):\n- GV: {step.teacher_action}\n- HS: {step.student_action}\n")
        row_cells[4].text = "\n".join(steps_text)

        # Định dạng kích thước font trong bảng
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9.5)

        if act.advanced_extension:
            p_adv = doc.add_paragraph()
            p_adv.paragraph_format.space_before = Pt(4)
            r_adv_title = p_adv.add_run("🌟 Thử thách nâng cao (HS khá giỏi): ")
            r_adv_title.font.bold = True
            r_adv_title.font.color.rgb = RGBColor(194, 65, 12)
            p_adv.add_run(act.advanced_extension)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # IV. GHI CHÚ PHÂN HÓA
    doc.add_heading("IV. GHI CHÚ PHÂN HÓA ĐỐI TƯỢNG HỌC SINH", level=2)
    doc.add_paragraph(plan.differentiation_notes)

    # Xuất ra BytesIO
    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream
